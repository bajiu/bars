# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches
import numpy as np;
import random,os







# 加法好朋友数
def freind(arg):
    # 可以说是垃圾方法了
    num = arg
    arr = []
    qa = np.random.randint( low = 1,high = 10, size = num )
    for i in qa:
        q = ''
        a = ''
        n = random.randint(1,9)
        m = random.randint(1,9)
        r = random.randint(0,3)
        # 只有加
        if r == 0:
            q = str(10 - n) + '+' + str(m) + '+' + str(n)
            a = 10 + m
        elif r == 1:
            q = str(10 - n) + '+' + str(n) + '+' + str(m)
            a = 10 + m
        # 还有减
        elif r == 2:
            q = str(10 - n) + '+' + str(n) + '-' + str(m)
            a = 10 - m
        else:
            q = str(10 - n) + '-' + str(m) + '+' + str(n)
            a = 10 - m
        arr.append((q,a))
    # print(arr)
    return arr


# 同尾数
def tail(arg):
    # 垃圾方法x2
    num = arg
    arr = []
    qa = np.random.randint( low = 1,high = 10, size = num )
    for i in qa:
        q = ''
        a = ''
        n = random.randint(1,9)
        r = random.randint(0,3)
        m = random.randint(1,9)
        # 只有加
        if r == 0:
            q = str(10 + n) + '+' + str(m) + '-' + str(n)
            a = 10 + m
        elif r == 1:
            q = str(10 + n) + '-' + str(n) + '+' + str(m)
            a = 10 + m
        # 还有减
        elif r == 2:
            q = str(10 + n) + '-' + str(n) + '-' + str(m)
            a = 10 - m
        else:
            q = str(10 + n) + '-' + str(m) + '-' + str(n)
            a = 10 - m
        arr.append((q,a))
    print(arr)
    return arr
# 带括号
def bracket(arg):
    num = arg
    arr = []
    qa = np.random.randint( low = 5,high = 10, size = num )
    for i in qa:
        o = random.randint(1,20)
        r = random.randint(0,1)
        m = random.randint(-20,20)
        k = i - o - m
        o = str(o)
        if k < 0:
            k = str(k)
        else:
            k = '+' + str(k)
        if r == 0:
            if m < 0:
                m = str(m)
            else:
                m = '+' + str(m)
            # 框框在前头
            o = '(' + o
            m = m + ')'
        else:
            # 框框在后头
            if m < 0:
                m = '-(' + str(abs(m))
            else:
                m = '+(' + str(m)
            k += ' )'
        q = o + m + k
        a = i
        arr.append((q,a))
    return arr
# 进位加法
def carryadd(arg):
    num = arg
    arr = []
    qa = np.random.randint( low = 5,high = 10, size = num )
    for i in qa:
        m = random.randint(5,10)
        q = str(i) + '+' + str(m)
        a = i + m
        arr.append((q,a))
    # print(arr)
    return arr
# 退位减法
def backsub(arg):
    num = arg
    arr = []
    qa = np.random.randint( low = 10,high = 20, size = num )
    for i in qa:
        m = random.randint((i-9),i)
        q = str(i) + '-' + str(m)
        a = i - m
        arr.append((q,a))
    print(arr)
    return arr

def pull(arg):
    num = arg
    arr = []
    # qa = np.random.randint( low = 1,high = 10, size = num )
    for i in range(num):
        q1 = random.randint(6,9)
        q2 = random.randint(6,9)
        q3 = random.randint(4,8)
        q = str(q1) + '+' + str(q2) + '+' + str(q3)
        a = q1 + q2 + q3
        arr.append((q,a))
    print(arr)
    return arr





def makeTopic(fileName = '' , title = '' ,type = '' , size = 20 , answer = False):
    folder = os.path.exists(os.path.abspath('小学一年级数学题'))
    mathPath = os.path.abspath('小学一年级数学题')
    if folder == False:
        os.makedirs(mathPath)


    t = ''
    if type == '':
        return
    elif type == 'freind':
        t = freind(size)
    elif type == 'tail':
        t = tail(size)
    elif type == 'carryadd':
        t = carryadd(size)
    elif type == 'backsub':
        t = backsub(size)
    elif type == 'bracket':
        t = bracket(size)
    elif type == 'pull':
        t = pull(size)

    document = Document()
    document.add_heading(title + '题目', 0)
    table = document.add_table(rows=2, cols=10)
    hdr_cells = table.rows[0].cells
    for i in range(10):
        hdr_cells[i].text = str(t[i][0])
    hdr_cells = table.rows[1].cells
    for j in range(10):
        hdr_cells[j].text = str(t[j+10][0])
    document.save(os.path.join(mathPath,fileName + '题目.docx'))
    if answer:
        document = Document()
        document.add_heading(title + '答案', 0)
        table = document.add_table(rows=2, cols=10)
        hdr_cells = table.rows[0].cells
        for i in range(10):
            hdr_cells[i].text = str(t[i][1])
        hdr_cells = table.rows[1].cells
        for j in range(10):
            hdr_cells[j].text = str(t[j+10][1])
        document.save(os.path.join(mathPath,fileName + '答案.docx'))





conf = [
    {
        'name': 'freind',
        'fileName': '加法好朋友数',
        'title': '加法好朋友数',
        'type': 'freind',
        'size': 20,
        'answer': False
    },
    {
       'name': 'tail',
       'fileName': '同尾数',
       'title': '同尾数',
       'type': 'tail',
       'size': 20,
       'answer': False
   },
   {
      'name': 'carryadd',
      'fileName': '进位加法',
      'title': '进位加法',
      'type': 'carryadd',
      'size': 20,
      'answer': False
  },
  {
     'name': 'backsub',
     'fileName': '退位减法',
     'title': '退位减法',
     'type': 'backsub',
     'size': 20,
     'answer': False
  },
  {
     'name' : 'bracket',
     'fileName': '带括号的随机加减法',
     'title': '带括号的随机加减法',
     'type': 'bracket',
     'size': 20,
     'answer': False
  },
  {
      'name': 'pull',
      'fileName': '拆小补大',
      'title': '拆小补大',
      'type': 'pull',
      'size': 20,
      'answer': False
  }
]

for i in conf:
    makeTopic(fileName = i['fileName'] , title = i['title'] , type = i['type'] , size = i['size'] , answer = i['answer'])


# makeTopic(fileName = conf[0]['fileName'] , title = conf[0]['title'] , type = conf[0]['type'] , size = conf[0]['size'] , answer = conf[0]['answer'])
